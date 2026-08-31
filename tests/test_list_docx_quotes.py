"""Quote inventory extraction from Word docs."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from list_docx_quotes import collect_quotes  # noqa: E402


def _make_fixture(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Introduction prose only.")
    doc.add_paragraph(
        'בית שמאי אומרים: בערב יטו ויקראו. (משנה ברכות א, ג)'
    )
    doc.add_paragraph(
        'אמרו לו: כדאי היית לחוב בעצמך. (משנה ברכות א, ג)'
    )
    doc.add_paragraph(
        'תנו רבנן: מעשה בחסיד אחד שהיה מתפלל בדרך. (בבלי ברכות לב ע"א)'
    )
    doc.save(str(path))


def test_collect_quotes_mishnah_and_bavli(tmp_path: Path) -> None:
    fx = tmp_path / "quotes.docx"
    _make_fixture(fx)
    quotes = collect_quotes(Document(str(fx)))
    assert len(quotes) >= 2
    corps = {q["corpus"] for q in quotes}
    assert "mishnah" in corps
    assert "bavli" in corps
    assert all(q["id"] == i + 1 for i, q in enumerate(quotes))
    assert all(q["ref"] for q in quotes)


def test_cli_json(tmp_path: Path) -> None:
    import json
    import subprocess

    fx = tmp_path / "quotes.docx"
    _make_fixture(fx)
    proc = subprocess.run(
        [
            sys.executable,
            str(ROOT / "scripts" / "list_docx_quotes.py"),
            "--input",
            str(fx),
            "--json",
        ],
        capture_output=True,
        text=True,
        check=True,
    )
    data = json.loads(proc.stdout)
    assert isinstance(data, list)
    assert len(data) >= 2
