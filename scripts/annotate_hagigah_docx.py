#!/usr/bin/env python3
"""Annotate Hagigah Aug 2025: short Hebrew RTL comments + tracked changes.

Includes Mishnah, Tosefta, Bavli. Links use ?misyzira=&mm15= (not FreeText ?query=).
"""

from __future__ import annotations

import argparse
import shutil
import sys
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
from maagarim_links import (  # noqa: E402
    BAVLI_HAG_MM15_2a,
    BAVLI_HAG_MM15_5b,
    TRACTATE_HAGIGAH,
    find_instructions,
    mm15_unit,
)

AUTHOR = "Maagarim Reader"
INITIALS = "MR"
ISO_DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

MESIRAH_MISHNAH = "Kaufmann A 50"
MESIRAH_TOSEFTA = "Wien 46"
MESIRAH_HAG_BAVLI = "Munich 6"
MESIRAH_RH_BAVLI = "JTS EMC 270"
ID_MISHNAH = 31000
ID_TOSEFTA = 28000
ID_HAG_BAVLI = 80023
ID_RH_BAVLI = 80019


class RevIds:
    def __init__(self, start: int = 1):
        self.n = start

    def next(self) -> int:
        i = self.n
        self.n += 1
        return i


def enable_track_revisions(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def set_paragraph_rtl(paragraph) -> None:
    """Mark a paragraph as RTL/bidi (for Hebrew comments)."""
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "right")
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        if rPr.find(qn("w:rtl")) is None:
            rPr.append(OxmlElement("w:rtl"))


def _run_text_len(run) -> int:
    return len(run.text or "")


def split_run_at(run, index: int):
    text = run.text or ""
    if index <= 0 or index >= len(text):
        return
    left, right = text[:index], text[index:]
    run.text = left
    new_el = deepcopy(run._r)
    t_els = new_el.findall(qn("w:t"))
    if t_els:
        t_els[0].text = right
        t_els[0].set(XML_SPACE, "preserve")
        for extra in t_els[1:]:
            extra.getparent().remove(extra)
    else:
        t = OxmlElement("w:t")
        t.set(XML_SPACE, "preserve")
        t.text = right
        new_el.append(t)
    run._r.addnext(new_el)


def split_paragraph_at(paragraph, char_index: int) -> None:
    if char_index <= 0:
        return
    acc = 0
    for run in list(paragraph.runs):
        n = _run_text_len(run)
        if acc + n < char_index:
            acc += n
            continue
        if acc + n == char_index:
            return
        split_run_at(run, char_index - acc)
        return


def runs_covering(paragraph, start: int, end: int):
    split_paragraph_at(paragraph, end)
    split_paragraph_at(paragraph, start)
    out = []
    acc = 0
    for run in paragraph.runs:
        n = _run_text_len(run)
        run_start, run_end = acc, acc + n
        acc = run_end
        if n == 0:
            continue
        if run_end <= start or run_start >= end:
            continue
        out.append(run)
    return out


def find_span(paragraph, needle: str):
    idx = paragraph.text.find(needle)
    if idx < 0:
        return None
    return idx, idx + len(needle)


def wrap_runs_in_one_del(runs, author: str, date: str, rev_id: int) -> etree._Element:
    els = [run._r for run in runs]
    parent = els[0].getparent()
    idx = parent.index(els[0])
    d = OxmlElement("w:del")
    d.set(qn("w:id"), str(rev_id))
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), date)
    for r in els:
        for t in r.findall(qn("w:t")):
            t.tag = qn("w:delText")
        parent.remove(r)
        d.append(r)
    parent.insert(idx, d)
    return d


def insert_ins_after(anchor_el, text: str, author: str, date: str, rev_id: int, rPr=None):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(deepcopy(rPr))
    t = OxmlElement("w:t")
    t.set(XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    anchor_el.addnext(ins)
    return ins


def tracked_replace(paragraph, needle: str, replacement: str, revs: RevIds) -> bool:
    span = find_span(paragraph, needle)
    if span is None:
        return False
    start, end = span
    covered = runs_covering(paragraph, start, end)
    if not covered:
        return False
    rPr = covered[0]._r.find(qn("w:rPr"))
    last_del = wrap_runs_in_one_del(covered, AUTHOR, ISO_DATE, revs.next())
    insert_ins_after(last_del, replacement, AUTHOR, ISO_DATE, revs.next(), rPr=rPr)
    return True


def add_comment_on_needle(doc: Document, paragraph, needle: str, text: str) -> bool:
    span = find_span(paragraph, needle)
    if span is None:
        runs = [r for r in paragraph.runs if r.text]
        if not runs:
            return False
        comment = doc.add_comment(runs, text=text, author=AUTHOR, initials=INITIALS)
    else:
        start, end = span
        covered = runs_covering(paragraph, start, end)
        if not covered:
            return False
        comment = doc.add_comment(covered, text=text, author=AUTHOR, initials=INITIALS)
    for p in comment.paragraphs:
        set_paragraph_rtl(p)
    return True


def comment_lines(*lines: str) -> str:
    return "\n".join(line for line in lines if line)


def short_find(
    composition: str,
    mesira: str,
    phrase: str,
    mishibbur: int,
    *,
    mm15: str | None = None,
    page: int | None = None,
) -> str:
    return find_instructions(
        composition=composition,
        mesira=mesira,
        search_phrase=phrase,
        mishibbur=mishibbur,
        mm15=mm15,
        page=page,
    )


def hag_mm(chapter: int, unit: int) -> str:
    """Hagigah mishnah/tosefta mm15 (tractate 023)."""
    return mm15_unit(TRACTATE_HAGIGAH, chapter, unit)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", "-i", type=Path, required=True)
    ap.add_argument("--output", "-o", type=Path, default=None)
    ap.add_argument("--backup", "-b", type=Path, required=True)
    args = ap.parse_args()

    original = args.input.expanduser().resolve()
    backup = args.backup.expanduser().resolve()
    if not backup.exists():
        raise SystemExit(f"Backup not found: {backup}")
    shutil.copy2(backup, original)

    doc = Document(str(original))
    enable_track_revisions(doc)
    revs = RevIds(1)
    paras = doc.paragraphs
    misses: list[str] = []

    def C(i: int, needle: str, text: str):
        if not add_comment_on_needle(doc, paras[i], needle, text):
            misses.append(f"comment miss p{i}: {needle[:40]!r}")

    def R(i: int, needle: str, replacement: str):
        if not tracked_replace(paras[i], needle, replacement, revs):
            misses.append(f"replace miss p{i}: {needle[:40]!r}")

    C(
        0,
        paras[0].text.strip() or paras[0].text,
        comment_lines(
            "בדיקת ציטוטי משנה / תוספתא / תלמוד במאגרים (GetYziraFull).",
            "משנה: Kaufmann A 50 · תוספתא: Wien 46 · בבלי חגיגה: Munich 6 · ר״ה: JTS EMC 270.",
            "Word → סקירה → כל הסימונים.",
        ),
    )

    # Live GetYziraFull comparisons (2026-08-27)
    C(
        12,
        "הכל חיבין בראיה",
        comment_lines(
            "שונה מ-Kaufmann A 50 (משנה חגיגה א, א).",
            "במאמר: חיבין | במסירה: חייבים.",
            short_find(
                "משנה", MESIRAH_MISHNAH, "הכל חייבים בראייה", ID_MISHNAH, mm15=hag_mm(1, 1)
            ),
        ),
    )
    C(
        13,
        "הראיה שתי כסף וחגיגה מעה כסף",
        comment_lines(
            "שונה מ-Kaufmann A 50 (משנה חגיגה א, ב).",
            "במאמר: הראיה | במסירה: הראייה.",
            short_find(
                "משנה", MESIRAH_MISHNAH, "הראייה שתי כסף", ID_MISHNAH, mm15=hag_mm(1, 2)
            ),
        ),
    )
    C(
        21,
        "יוחנן בן דהבאי אומר משום רבי יהודה",
        comment_lines(
            "שונה מ-Munich 6 (בבלי חגיגה ב ע\"א).",
            "במאמר: אומר…רבי…הראייה | במסירה: או'…ר'…הראיה.",
            short_find(
                "תלמוד בבלי, חגיגה",
                MESIRAH_HAG_BAVLI,
                "יוחנן בן דהבאי",
                ID_HAG_BAVLI,
                mm15=BAVLI_HAG_MM15_2a,
            ),
        ),
    )
    C(
        26,
        "אין דורשין בעריות בשלשה",
        comment_lines(
            "שונה מ-Kaufmann A 50 (משנה חגיגה ב, א).",
            "במאמר: דורשין…בשלשה | במסירה: דורשים…בשלושה.",
            short_find(
                "משנה", MESIRAH_MISHNAH, "אין דורשים בעריות", ID_MISHNAH, mm15=hag_mm(2, 1)
            ),
        ),
    )
    C(
        30,
        "ארבעה נכנסו לפרדס",
        comment_lines(
            "שונה מ-Wien 46 (תוספתא חגיגה ב, ג).",
            "במאמר: בן עזאי…קצץ | במסירה: בן עזיי…קיצץ.",
            short_find(
                "תוספתא",
                MESIRAH_TOSEFTA,
                "ארבעה נכנסו לפרדס",
                ID_TOSEFTA,
                mm15=hag_mm(2, 3),
            ),
        ),
    )
    C(
        32,
        "לפרדס של מלך ועלייה בנויה",
        comment_lines(
            "תואם Wien 46 (תוספתא חגיגה ב, ה).",
            short_find(
                "תוספתא", MESIRAH_TOSEFTA, "לפרדס של מלך", ID_TOSEFTA, mm15=hag_mm(2, 5)
            ),
        ),
    )
    C(
        36,
        "חייב אדם לטהר את עצמו ברגל",
        comment_lines(
            "שונה מ-JTS EMC 270 (בבלי ר״ה טז ע\"ב).",
            "במאמר: לטהר את עצמו | במסירה: לטהר עצמו.",
            short_find(
                "תלמוד בבלי, ראש השנה", MESIRAH_RH_BAVLI, "חייב אדם לטהר", ID_RH_BAVLI
            ),
        ),
    )
    C(
        37,
        "משעבר הרגל היו מעבירין על טהרת עזרה",
        comment_lines(
            "שונה מ-Kaufmann A 50 (משנה חגיגה ג, ז).",
            "במאמר: טהרת עזרה | במסירה: טהרת העזרה.",
            short_find(
                "משנה", MESIRAH_MISHNAH, "משעבר הרגל", ID_MISHNAH, mm15=hag_mm(3, 7)
            ),
        ),
    )
    C(
        39,
        "טעונין טבילה, חוץ ממזבח הזהב",
        comment_lines(
            "שונה מ-Kaufmann A 50 (משנה חגיגה ג, ח).",
            "במאמר: טעונין…הנחושת…מצופין | במסירה: טעונים…הנחשת…מצופים.",
            short_find(
                "משנה", MESIRAH_MISHNAH, "טעונים טבילה", ID_MISHNAH, mm15=hag_mm(3, 8)
            ),
        ),
    )
    C(
        44,
        "מטבילין מאור הלבנה",
        comment_lines(
            "תואם Wien 46 (תוספתא חגיגה, מאור הלבנה).",
            short_find(
                "תוספתא",
                MESIRAH_TOSEFTA,
                "מאור הלבנה",
                ID_TOSEFTA,
                mm15=hag_mm(3, 25),
            ),
        ),
    )
    C(
        49,
        "היתר נדרים פורחין באויר",
        comment_lines(
            "שונה מ-Kaufmann A 50 (משנה חגיגה א, ח).",
            "במאמר: היתר…פורחין…שיסמוכו | במסירה: התר…פורחים…שיסמכו.",
            short_find(
                "משנה", MESIRAH_MISHNAH, "התר נדרים פורחים", ID_MISHNAH, mm15=hag_mm(1, 8)
            ),
        ),
    )
    C(
        55,
        "יוסף בן יועזר אומר שלא לסמוך",
        comment_lines(
            "שונה מ-Kaufmann A 50 (משנה חגיגה ב, ב).",
            "במאמר: יוסף | במסירה: יוסה.",
            short_find(
                "משנה", MESIRAH_MISHNAH, "יוסה בן יועזר", ID_MISHNAH, mm15=hag_mm(2, 2)
            ),
        ),
    )
    C(
        56,
        "מביאין שלמים ואין סומכין עליהן",
        comment_lines(
            "שונה מ-Kaufmann A 50 (משנה חגיגה ב, ג).",
            "במאמר (שמאי): עליהן | במסירה: עליהם.",
            short_find(
                "משנה", MESIRAH_MISHNAH, "מביאין שלמים", ID_MISHNAH, mm15=hag_mm(2, 3)
            ),
        ),
    )
    C(
        58,
        "בתחילה לא היתה מחלוקת בישראל",
        comment_lines(
            "שונה מ-Wien 46 (תוספתא חגיגה ב, ט).",
            "במאמר: בתחילה…ר' יוסי | במסירה: כתחלה…ר' יהודה.",
            short_find(
                "תוספתא",
                MESIRAH_TOSEFTA,
                "לא היתה מחלוקת",
                ID_TOSEFTA,
                mm15=hag_mm(2, 9),
            ),
        ),
    )
    C(
        65,
        "רבי ורבי חייא הוו שקלי ואזלי באורחא",
        comment_lines(
            "שונה מ-Munich 6 (בבלי חגיגה ה ע\"ב) — כל הקטע.",
            "במאמר: ורבי…שקלי ואזלי…לההוא…נזיל וניקביל…אמרי:",
            "במסירה: ור'…קאזלי…ההיא…ניעול נקביל…אמרי להו.",
            short_find(
                "תלמוד בבלי, חגיגה",
                MESIRAH_HAG_BAVLI,
                "קאזלי באורחא",
                ID_HAG_BAVLI,
                mm15=BAVLI_HAG_MM15_5b,
            ),
        ),
    )
    C(
        66,
        "אמר ליה ר' חייא לרבי",
        comment_lines(
            "המשך Munich 6 (ה ע\"ב).",
            "במאמר: אמר…תיב את, לא…איזיל אנא ואקביל אפיה.",
            "במסירה: אמ'…תיב את דלא…איעול אנא אקביל אפי קמיה ואיתי.",
            short_find(
                "תלמוד בבלי, חגיגה",
                MESIRAH_HAG_BAVLI,
                "קאזלי באורחא",
                ID_HAG_BAVLI,
                mm15=BAVLI_HAG_MM15_5b,
            ),
        ),
    )

    # Suggestion-mode tracked changes
    R(12, "הכל חיבין בראיה", "הכל חייבים בראייה")
    R(13, "הראיה שתי כסף וחגיגה מעה כסף", "הראייה שתי כסף וחגיגה מעה כסף")
    R(13, "הראיה מעה כסף", "הראייה מעה כסף")
    R(21, "אומר משום רבי יהודה", "או' משום ר' יהודה")
    R(21, "מן הראייה", "מן הראיה")
    R(26, "אין דורשין בעריות בשלשה", "אין דורשים בעריות בשלושה")
    R(30, "בן עזאי ובן זומא", "בן עזיי ובן זומא")
    R(30, "אחד הציץ וקצץ בנטיעות", "אחד הציץ וקיצץ בנטיעות")
    R(36, "לטהר את עצמו ברגל", "לטהר עצמו ברגל")
    R(37, "טהרת עזרה", "טהרת העזרה")
    R(39, "טעונין טבילה", "טעונים טבילה")
    R(39, "מזבח הנחושת", "מזבח הנחשת")
    R(39, "מפני שהן מצופין", "מפני שהן מצופים")
    R(49, "היתר נדרים פורחין באויר", "התר נדרים פורחים באויר")
    R(49, "שיסמוכו", "שיסמכו")
    R(55, "יוסף בן יועזר", "יוסה בן יועזר")
    R(55, "יוסף בן יוחנן", "יוסה בן יוחנן")
    R(56, "ואין סומכין עליהן, אבל לא עולות", "ואין סומכין עליהם, אבל לא עולות")
    R(58, "בתחילה לא היתה מחלוקת", "כתחלה לא היתה מחלוקת")
    # Bavli Hagigah 5b — full quoted segment vs Munich
    R(65, "רבי ורבי חייא", "רבי ור' חייא")
    R(65, "הוו שקלי ואזלי באורחא", "הוו קאזלי באורחא")
    R(65, "כי מטו לההוא מתא", "כי מטו ההיא מתא")
    R(65, "נזיל וניקביל אפיה", "ניעול נקביל אפיה")
    R(65, "אמרי: איכא צורבא מרבנן הכא, ומאור עינים הוא", "אמרי להו. איכא צורבא מרבנן הכא ומאור עינים הוא")
    R(66, "אמר ליה ר' חייא לרבי: תיב את, לא תזלזל בנשיאותך. איזיל אנא ואקביל אפיה.", "אמ' ליה ר' חייא לרבי. תיב את דלא תזלזל בנשיאותך. איעול אנא אקביל אפי קמיה ואיתי.")

    doc.save(str(original))
    out = args.output or (ROOT / "output" / "hagigah-aug-2025-first-witness.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(original, out)

    z = Document(str(original))
    xml = z.element.body.xml
    print("saved", original)
    print("comments", len(list(z.comments)), "ins", xml.count("<w:ins "), "del", xml.count("<w:del "))
    # sample
    c0 = list(z.comments)[1]
    print("sample:", " | ".join(p.text for p in c0.paragraphs)[:200])
    for m in misses:
        print("MISS", m)


if __name__ == "__main__":
    main()
