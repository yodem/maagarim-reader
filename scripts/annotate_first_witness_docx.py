#!/usr/bin/env python3
"""Annotate a Word doc: Maagarim first-witness check with comments + tracked changes.

Only Mishnah and Talmud verbatim quotes are annotated. Tanakh, ketubah formulas,
and paraphrases are skipped (no comment).
"""

from __future__ import annotations

import argparse
import shutil
import sys
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
from maagarim_links import find_instructions, search_link  # noqa: E402

DEFAULT_BACKUP = ROOT / "examples" / "fixtures" / "ketubot-source.docx"
FALLBACK_BACKUP = ROOT / "output" / "ketubot-nov-2025-source.docx"
DEFAULT_OUTPUT = ROOT / "output" / "annotated.docx"

AUTHOR = "Maagarim Reader"
INITIALS = "MR"
ISO_DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

MESIRAH_MISHNAH = "Kaufmann A 50"
MESIRAH_KET_BAVLI = "Vatican, Biblioteca Apostolica ebr., 130"
MESIRAH_YEV_BAVLI = "העדות הראשונה במאגרים ל«תלמוד בבלי, יבמות»"
MESIRAH_KID_BAVLI = "העדות הראשונה במאגרים ל«תלמוד בבלי, קידושין»"
MESIRAH_YER_KET = "העדות הראשונה במאגרים ל«תלמוד ירושלמי, כתובות»"

V_2A_SHAKDU_DIPL = (
    "שקדו חכמ' שתינ[ש?א?] ברביעי בשבת שיהא טורח בסעודה [שלשה ימים]. "
    "אחד בשבת ושיני בשבת ושלישי בשבת. וברביעי כונסה."
)
V_2A_SHAKDU_READ = (
    "שקדו חכמ' שתינשא ברביעי בשבת שיהא טורח בסעודה שלשה ימים. "
    "אחד בשבת ושיני בשבת ושלישי בשבת. וברביעי כונסה."
)

K_1_1_DIPL = (
    "בתולה נישאת ביום הרביעי ואלמנה ביום החמישי. "
    "שפעמים בשבת בתי דינין יושבין בעיירות ביום השיני וביום החמישי. "
    "שאם היה ל(י)[ו] טענת בתולים היה משכים לבית דין."
)
K_1_1_READ = (
    "בתולה נישאת ביום הרביעי ואלמנה ביום החמישי. "
    "שפעמים בשבת בתי דינין יושבין בעיירות ביום השיני וביום החמישי. "
    "שאם היה לו טענת בתולים היה משכים לבית דין."
)
K_9_4_READ = (
    "המושיב אשתו חנוונית או שמינה אפיטרופוס. "
    "הרי זה משביעה כל זמן שירצה. ר' אליעזר אומ'. אפילו על פילכה ועל עיסתה."
)
K_9_4_DIPL = K_9_4_READ
K_2_2 = "שהפה שאסר הוא הפה שהיתיר"
K_2_5A = "האשה שאמרה. אשת איש היתי וגרושה אני. נאמנת. שהפה שאסר הוא הפה שהיתיר."
K_2_5B = "האשה שאמרה. נשביתי וטהורה אני. נאמנת. שהפה שאסר הוא הפה שהיתיר."
K_9_1_END = "שכל המתנה על הכתוב שבתורה תנאו בטל"


class RevIds:
    def __init__(self, start: int = 1):
        self.n = start

    def next(self) -> int:
        i = self.n
        self.n += 1
        return i


def enable_track_revisions(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def _run_text_len(run) -> int:
    return len(run.text or "")


def split_run_at(run, index: int):
    text = run.text or ""
    if index <= 0 or index >= len(text):
        return
    left, right = text[:index], text[index:]
    run.text = left
    new_el = deepcopy(run._r)
    t_els = new_el.findall(qn("w:t"))
    if t_els:
        t_els[0].text = right
        t_els[0].set(XML_SPACE, "preserve")
        for extra in t_els[1:]:
            extra.getparent().remove(extra)
    else:
        t = OxmlElement("w:t")
        t.set(XML_SPACE, "preserve")
        t.text = right
        new_el.append(t)
    run._r.addnext(new_el)


def split_paragraph_at(paragraph, char_index: int) -> None:
    if char_index <= 0:
        return
    acc = 0
    for run in list(paragraph.runs):
        n = _run_text_len(run)
        if acc + n < char_index:
            acc += n
            continue
        if acc + n == char_index:
            return
        split_run_at(run, char_index - acc)
        return


def runs_covering(paragraph, start: int, end: int):
    split_paragraph_at(paragraph, end)
    split_paragraph_at(paragraph, start)
    out = []
    acc = 0
    for run in paragraph.runs:
        n = _run_text_len(run)
        run_start, run_end = acc, acc + n
        acc = run_end
        if n == 0:
            continue
        if run_end <= start or run_start >= end:
            continue
        out.append(run)
    return out


def find_span(paragraph, needle: str):
    text = paragraph.text
    idx = text.find(needle)
    if idx < 0:
        return None
    return idx, idx + len(needle)


def wrap_runs_in_one_del(runs, author: str, date: str, rev_id: int) -> etree._Element:
    els = [run._r for run in runs]
    parent = els[0].getparent()
    idx = parent.index(els[0])
    d = OxmlElement("w:del")
    d.set(qn("w:id"), str(rev_id))
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), date)
    for r in els:
        for t in r.findall(qn("w:t")):
            t.tag = qn("w:delText")
        parent.remove(r)
        d.append(r)
    parent.insert(idx, d)
    return d


def insert_ins_after(anchor_el, text: str, author: str, date: str, rev_id: int, rPr=None):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(deepcopy(rPr))
    t = OxmlElement("w:t")
    t.set(XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    anchor_el.addnext(ins)
    return ins


def tracked_replace(paragraph, needle: str, replacement: str, revs: RevIds) -> bool:
    span = find_span(paragraph, needle)
    if span is None:
        return False
    start, end = span
    covered = runs_covering(paragraph, start, end)
    if not covered:
        return False
    rPr = covered[0]._r.find(qn("w:rPr"))
    last_del = wrap_runs_in_one_del(covered, AUTHOR, ISO_DATE, revs.next())
    insert_ins_after(last_del, replacement, AUTHOR, ISO_DATE, revs.next(), rPr=rPr)
    return True


def add_comment_on_needle(doc: Document, paragraph, needle: str, text: str) -> bool:
    span = find_span(paragraph, needle)
    if span is None:
        runs = [r for r in paragraph.runs if r.text]
        if not runs:
            return False
        doc.add_comment(runs, text=text, author=AUTHOR, initials=INITIALS)
        return True
    start, end = span
    covered = runs_covering(paragraph, start, end)
    if not covered:
        return False
    doc.add_comment(covered, text=text, author=AUTHOR, initials=INITIALS)
    return True


def comment_lines(*lines: str) -> str:
    return "\n".join(line for line in lines if line is not None)


def mishnah_find(search_phrase: str, *, cite: str | None = None) -> str:
    return find_instructions(
        composition="משנה",
        mesira=MESIRAH_MISHNAH,
        search_phrase=search_phrase,
        cite=cite,
    )


def bavli_ket_find(search_phrase: str, *, cite: str | None = None) -> str:
    return find_instructions(
        composition="תלמוד בבלי, כתובות",
        mesira=MESIRAH_KET_BAVLI,
        search_phrase=search_phrase,
        cite=cite,
    )


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Ketubot reference run: Maagarim first-witness comments + tracked changes."
    )
    p.add_argument("--input", "-i", type=Path, required=True, help="Source .docx to annotate")
    p.add_argument(
        "--output",
        "-o",
        type=Path,
        default=None,
        help="Optional copy of annotated file (default: annotate --input in place)",
    )
    p.add_argument(
        "--backup",
        "-b",
        type=Path,
        default=None,
        help="Clean backup restored before each run (default: examples/fixtures/ketubot-source.docx)",
    )
    return p.parse_args()


def resolve_backup(explicit: Path | None) -> Path:
    if explicit is not None:
        return explicit
    if DEFAULT_BACKUP.exists():
        return DEFAULT_BACKUP
    if FALLBACK_BACKUP.exists():
        return FALLBACK_BACKUP
    raise SystemExit(
        "No clean backup found. Pass --backup path/to/clean.docx or add "
        "examples/fixtures/ketubot-source.docx"
    )


def main() -> None:
    args = parse_args()
    original = args.input.expanduser().resolve()
    if not original.exists():
        raise SystemExit(f"Source not found: {original}")

    backup = resolve_backup(args.backup)
    if not backup.exists():
        shutil.copy2(original, backup)
    shutil.copy2(backup, original)
    doc = Document(str(original))
    enable_track_revisions(doc)
    revs = RevIds(1)
    paras = doc.paragraphs
    misses: list[str] = []

    def C(i: int, needle: str, text: str):
        ok = add_comment_on_needle(doc, paras[i], needle, text)
        if not ok:
            misses.append(f"comment miss p{i}: {needle[:40]!r}")

    def R(i: int, needle: str, replacement: str):
        ok = tracked_replace(paras[i], needle, replacement, revs)
        if not ok:
            misses.append(f"replace miss p{i}: {needle[:40]!r}")

    C(
        0,
        paras[0].text.strip() or paras[0].text,
        comment_lines(
            "מדריך קריאה — בדיקת ציטוטי משנה ותלמוד במאגרים (27.8.2026)",
            "",
            "מה נעשה?",
            "השווינו ציטוטי משנה ותלמוד במאמר מול הנוסח במאגרים של האקדמיה ללשון העברית.",
            "אם לא ציינתם כתב יד — השתמשנו בעדות הראשונה במאגרים לכל חיבור:",
            "• משנה → Kaufmann A 50",
            "• בבלי כתובות → Vatican ebr. 130",
            "• בבלי יבמות / קידושין, ירושלמי כתובות → העדות הראשונה במאגרים",
            "",
            "איך לקרוא?",
            "Word → סקירה → «כל הסימונים». מחיקה (אדום) = הנוסח במאמר; הוספה = הנוסח במאגרים.",
            "הערות בשוליים מסבירות את ההבדל ומראות איך למצוא את המקור.",
            "",
            "אתר: maagarim.hebrew-academy.org.il",
        ),
    )

    q11 = (
        "בתולה נשאת ליום הרביעי, ואלמנה ליום החמישי; "
        "שפעמים בשבת בתי דינין יושבין בעיירות, ביום השני וביום החמישי, "
        "שאם היה לו טענת בתולים היה משכים לבית דין."
    )
    q94 = (
        "המושיב את אשתו חנונית או שמינה אותה אפוטרופא [=אפוטרופוס]—"
        "הרי זה משביעה [שלא לקחה לעצמה] כל זמן שירצה. "
        "רבי אליעזר אומר: אפילו על פלכה ועל עיסתה."
    )
    q16 = "לא מפיה אנו חיין... עד שתביא ראיה לדבריה"
    q91 = "כל המתנה על מה שכתוב בתורה תנאו בטל"
    q_shakdu = (
        "שקדו חכמים על תקנת בנות ישראל שיהא טורח בסעודה שלשה ימים, "
        "אחד בשבת ושני בשבת ושלישי בשבת, וברביעי כונס"
    )

    C(
        75,
        q11,
        comment_lines(
            "הנוסח במאמר שונה מ-Kaufmann A 50 (משנה כתובות א, א).",
            "במאמר: נשאת / ליום / השני.",
            "ב-Kaufmann: נישאת / ביום / השיני.",
            "נוסח המסירה: " + K_1_1_DIPL,
            "",
            mishnah_find("בתולה נשאת ליום הרביעי", cite="משנה כתובות א, א."),
        ),
    )
    C(
        57,
        q94,
        comment_lines(
            "הנוסח במאמר שונה מ-Kaufmann A 50 (משנה כתובות ט, ד).",
            "ב-Kaufmann: המושיב אשתו; חנוונית; אפיטרופוס; פילכה; ר' אליעזר אומ'.",
            "הסוגריים [=אפוטרופוס] ו-[שלא לקחה לעצמה] — הערות עורך, לא במסירה.",
            "נוסח Kaufmann: " + K_9_4_DIPL,
            "",
            mishnah_find("המושיב אשתו חנוונית", cite="משנה כתובות ט, ד."),
        ),
    )
    C(
        55,
        q16,
        comment_lines(
            "הנוסח במאמר שונה מ-Kaufmann A 50 (משנה כתובות א, ו–ט).",
            "במאמר: חיין / ראיה. ב-Kaufmann: חיים / ראייה.",
            "",
            mishnah_find("לא מפיה אנו חיים", cite="משנה כתובות א, ו–ט."),
        ),
    )
    C(
        55,
        "הפה שאסר הוא הפה שהתיר",
        comment_lines(
            "הנוסח במאמר שונה מ-Kaufmann A 50 (משנה כתובות ב, ב).",
            "ב-Kaufmann: «שהפה שאסר הוא הפה שהיתיר».",
            "",
            mishnah_find("שהפה שאסר הוא הפה שהיתיר", cite="משנה כתובות ב, ב."),
        ),
    )
    C(
        55,
        "האשה שאמרה: אשת איש הייתי וגרושה אני—נאמנת, שהפה שאסר",
        comment_lines(
            "הנוסח במאמר שונה מ-Kaufmann A 50 (משנה כתובות ב, ה, סעיף ראשון).",
            "ב-Kaufmann: «" + K_2_5A + "»",
            "",
            mishnah_find("אשת איש היתי וגרושה", cite="משנה כתובות ב, ה."),
        ),
    )
    C(
        55,
        "אמרה נשביתי וטהורה אני—נאמנת, שהפה שאסר הוא הפה שהתיר",
        comment_lines(
            "הנוסח במאמר שונה מ-Kaufmann A 50 (משנה כתובות ב, ה, סעיף שני).",
            "ב-Kaufmann: «" + K_2_5B + "»",
            "",
            mishnah_find("נשביתי וטהורה אני", cite="משנה כתובות ב, ה."),
        ),
    )

    C(
        78,
        "בתולה נישאת ביום הרביעי",
        comment_lines(
            "הפתיחה «בתולה נישאת ביום» תואמת את Kaufmann A 50 במשנה א, א —",
            "בניגוד לנוסח המודפס שצוטט קודם במאמר (נשאת / ליום).",
            "שאר הפסקה — ירושלמי, לא משנה Kaufmann.",
            "",
            mishnah_find("בתולה נישאת ביום הרביעי", cite="משנה כתובות א, א."),
        ),
    )

    C(
        45,
        q91,
        comment_lines(
            "הנוסח במאמר שונה מ-Kaufmann A 50 (משנה כתובות ט, א).",
            "במאמר: «על מה שכתוב בתורה».",
            "ב-Kaufmann: «שכל המתנה על הכתוב שבתורה תנאו בטל».",
            "",
            mishnah_find("כל המתנה על הכתוב שבתורה", cite="משנה כתובות ט, א."),
        ),
    )

    # Bavli Ketubot quotes — only comment on Talmud sources
    C(
        44,
        "תיקנו מזונותיה תחת מעשה ידיה, וקבורתה תחת כתובתה",
        comment_lines(
            "ציטוט מבבלי כתובות מז ע\"ב.",
            f"נבדק מול {MESIRAH_KET_BAVLI}.",
            "בבדיקה זו לא הושוו מילים במסירה (הציטוט מדף אחר).",
            "",
            bavli_ket_find("תיקנו מזונותיה תחת מעשה ידיה", cite="בבלי כתובות מז ע\"ב."),
        ),
    )
    C(
        59,
        "אם אומר את[ה] כן אין שלום בתוך ביתו לעולם",
        comment_lines(
            "ציטוט מירושלמי כתובות.",
            f"עדות: {MESIRAH_YER_KET}.",
            "בבדיקה זו לא הושוו מילים במסירה.",
            "",
            find_instructions(
                composition="תלמוד ירושלמי, כתובות",
                mesira=MESIRAH_YER_KET,
                search_phrase="אין שלום בתוך ביתו לעולם",
                cite="ירושלמי כתובות.",
            ),
        ),
    )
    C(
        60,
        "רבי אליעזר אומר: אף על פי שלא הושיבה חנוונית",
        comment_lines(
            "ציטוט מבבלי כתובות (ברייתא, דף מו ע\"ב).",
            f"עדות: {MESIRAH_KET_BAVLI}.",
            "בבדיקה זו לא הושוו מילים במסירה.",
            "",
            bavli_ket_find("אף על פי שלא הושיבה חנוונית", cite="בבלי כתובות."),
        ),
    )
    C(
        65,
        "נשא אדם אשה בילדותו, ישא אשה בזקנותו",
        comment_lines(
            "ציטוט מבבלי יבמotes (סב ע\"ב).",
            f"עדות: {MESIRAH_YEV_BAVLI}.",
            "בבדיקה זו לא הושוו מילים במסירה.",
            "",
            find_instructions(
                composition="תלמוד בבלי, יבמות",
                mesira=MESIRAH_YEV_BAVLI,
                search_phrase="נשא אדם אשה בילדותו",
                cite="בבלי יבמות.",
            ),
        ),
    )
    C(
        66,
        "נשא אשה ושהה עמה עשר שנים ולא ילדה אינו רשאי ליבטל",
        comment_lines(
            "ציטוט ממשנה יבמות (ו, ח) — באותה מסירת Kaufmann A 50 כמו המשנה במאגרים.",
            "בבדיקה זו לא הושוו מילים במסירה.",
            "",
            mishnah_find("נשא אשה ושהה עמה עשר שנים", cite="משנה יבמות ו, ח."),
        ),
    )
    C(
        68,
        "בן עשרים שנה ולא נשא אשה",
        comment_lines(
            "ציטוט מבבלי קידושין (כט ע\"ב).",
            f"עדות: {MESIRAH_KID_BAVLI}.",
            "בבדיקה זו לא הושוו מילים במסירה.",
            "",
            find_instructions(
                composition="תלמוד בבלי, קידושין",
                mesira=MESIRAH_KID_BAVLI,
                search_phrase="בן עשרים שנה ולא נשא",
                cite="בבלי קידושין.",
            ),
        ),
    )
    C(
        69,
        "דיינו שמגדלות בנינו ומצילות אותנו מן החטא",
        comment_lines(
            "ציטוט מבבלי יבמות (סג ע\"א–ע\"ב).",
            f"עדות: {MESIRAH_YEV_BAVLI}.",
            "בבדיקה זו לא הושוו מילים במסירה.",
            "",
            find_instructions(
                composition="תלמוד בבלי, יבמות",
                mesira=MESIRAH_YEV_BAVLI,
                search_phrase="מגדלות בנינו ומצילות",
                cite="בבלי יבמות.",
            ),
        ),
    )
    C(
        71,
        "מפריה ורביה בטיל",
        comment_lines(
            "ציטוט מבבלי יבמות (סא ע\"ב).",
            f"עדות: {MESIRAH_YEV_BAVLI}.",
            "בבדיקה זו לא הושוו מילים במסירה.",
            "",
            find_instructions(
                composition="תלמוד בבלי, יבמות",
                mesira=MESIRAH_YEV_BAVLI,
                search_phrase="מפריה ורביה בטיל",
                cite="בבלי יבמות.",
            ),
        ),
    )
    C(
        81,
        q_shakdu,
        comment_lines(
            "הנוסח במאמר שונה מ-Vatican ebr. 130 (בבלי כתובות ב ע\"א).",
            "במאמר: «שקדו חכמים על תקנת בנות ישראל… ושני… כונס».",
            "ב-Vatican: " + V_2A_SHAKDU_DIPL,
            "«על תקנת בנות ישראל» — נוסח דפוס, לא בפסקה זו ב-Vatican.",
            "",
            bavli_ket_find("שקדו חכמים שתינשא", cite="בבלי כתובות ב ע\"א."),
        ),
    )
    C(
        81,
        'שקדו חכמים על תקנת בנות ישראל", כלומר',
        comment_lines(
            "«שקדו חכמים על תקנת בנות ישראל» — נוסח דפוס.",
            "ב-Vatican ebr. 130 (דף ב ע\"א) הפסקה נפתחת אחרת.",
            "",
            "קישור לחיפוש: " + search_link("שקדו חכמים שתינשא"),
        ),
    )

    R(75, q11, K_1_1_READ)
    R(57, q94, K_9_4_READ)
    R(55, q16, "לא מפיה אנו חיים... עד שתביא ראייה לדבריה")
    R(55, "הפה שאסר הוא הפה שהתיר", K_2_2)
    R(55, "אשת איש הייתי", "אשת איש היתי")
    R(55, "הפה שהתיר", "הפה שהיתיר")
    R(55, "הפה שהתיר", "הפה שהיתיר")
    R(45, q91, K_9_1_END)
    R(81, q_shakdu, V_2A_SHAKDU_READ)

    doc.save(str(original))
    out_copy = args.output or DEFAULT_OUTPUT
    out_copy.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(original, out_copy)

    zdoc = Document(str(original))
    n_comments = len(list(zdoc.comments))
    xml = zdoc.element.body.xml
    n_ins = xml.count("<w:ins ")
    n_del = xml.count("<w:del ")
    print("saved", original)
    print("copy", out_copy)
    print("comments", n_comments, "ins", n_ins, "del", n_del)
    for m in misses:
        print("MISS", m)


if __name__ == "__main__":
    main()
