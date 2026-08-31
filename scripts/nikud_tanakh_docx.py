#!/usr/bin/env python3
"""Add nikud (without te'amim) to unpointed Tanakh quotations in a Word doc.

Flow (mirrors Sefaria gdocs Find Citations + sefaria-mcp get_text):
  1. Extract document text
  2. POST /api/find-refs (he + en) → Tanakh citations
  3. GET /api/texts/…?vhe=Tanach with Nikkud for each ref
  4. Locate matching unpointed (or trop-bearing) Hebrew spans
  5. Replace with menukadim bli te'amim via Word tracked changes

Usage:
  python3 scripts/nikud_tanakh_docx.py --input article.docx
  python3 scripts/nikud_tanakh_docx.py --input article.docx --dry-run
  python3 scripts/nikud_tanakh_docx.py --input article.docx --output other.docx

By default writes **back to the same file** with track revisions (Word suggestion mode).
A one-time backup is saved alongside as ``<name>-pre-nikud-backup.docx`` unless ``--no-backup``.
"""

from __future__ import annotations

import argparse
import shutil
import sys
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from hebrew_nikud import (  # noqa: E402
    letters_only,
    match_unpointed_in_haystack,
    menukad_bli_teamim,
    needs_nikud,
    needs_teamim_strip,
)
from sefaria_client import (  # noqa: E402
    find_refs,
    get_text_nikud,
    tanakh_refs_from_find_result,
)

AUTHOR = "Tanakh Nikud"
INITIALS = "TN"
ISO_DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
DEFAULT_OUTPUT = ROOT / "output" / "nikud-tanakh.docx"


def backup_path_for(input_path: Path) -> Path:
    return input_path.with_name(f"{input_path.stem}-pre-nikud-backup{input_path.suffix}")

# Chunk size for find-refs on long docs (linker NER).
CHUNK_CHARS = 6000
CHUNK_OVERLAP = 400


class RevIds:
    def __init__(self, start: int = 1):
        self.n = start

    def next(self) -> int:
        i = self.n
        self.n += 1
        return i


def enable_track_revisions(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def _run_text_len(run) -> int:
    return len(run.text or "")


def split_run_at(run, index: int) -> None:
    text = run.text or ""
    if index <= 0 or index >= len(text):
        return
    left, right = text[:index], text[index:]
    run.text = left
    new_el = deepcopy(run._r)
    t_els = new_el.findall(qn("w:t"))
    if t_els:
        t_els[0].text = right
        t_els[0].set(XML_SPACE, "preserve")
        for extra in t_els[1:]:
            extra.getparent().remove(extra)
    else:
        t = OxmlElement("w:t")
        t.set(XML_SPACE, "preserve")
        t.text = right
        new_el.append(t)
    run._r.addnext(new_el)


def split_paragraph_at(paragraph, char_index: int) -> None:
    if char_index <= 0:
        return
    acc = 0
    for run in list(paragraph.runs):
        n = _run_text_len(run)
        if acc + n < char_index:
            acc += n
            continue
        if acc + n == char_index:
            return
        split_run_at(run, char_index - acc)
        return


def runs_covering(paragraph, start: int, end: int):
    split_paragraph_at(paragraph, end)
    split_paragraph_at(paragraph, start)
    out = []
    acc = 0
    for run in paragraph.runs:
        n = _run_text_len(run)
        run_start, run_end = acc, acc + n
        acc = run_end
        if n == 0:
            continue
        if run_end <= start or run_start >= end:
            continue
        out.append(run)
    return out


def wrap_runs_in_one_del(runs, author: str, date: str, rev_id: int) -> etree._Element:
    els = [run._r for run in runs]
    parent = els[0].getparent()
    idx = parent.index(els[0])
    d = OxmlElement("w:del")
    d.set(qn("w:id"), str(rev_id))
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), date)
    for r in els:
        for t in r.findall(qn("w:t")):
            t.tag = qn("w:delText")
        parent.remove(r)
        d.append(r)
    parent.insert(idx, d)
    return d


def insert_ins_after(anchor_el, text: str, author: str, date: str, rev_id: int, rPr=None):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(deepcopy(rPr))
    t = OxmlElement("w:t")
    t.set(XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    anchor_el.addnext(ins)
    return ins


def tracked_replace_span(
    paragraph, start: int, end: int, replacement: str, revs: RevIds
) -> bool:
    covered = runs_covering(paragraph, start, end)
    if not covered:
        return False
    rPr = covered[0]._r.find(qn("w:rPr"))
    last_del = wrap_runs_in_one_del(covered, AUTHOR, ISO_DATE, revs.next())
    insert_ins_after(last_del, replacement, AUTHOR, ISO_DATE, revs.next(), rPr=rPr)
    return True


def iter_body_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def build_corpus(doc: Document) -> tuple[str, list[tuple[int, int, object]]]:
    """Concatenate paragraphs with '\\n' separators; map global offsets → para.

    Returns (corpus, spans) where each span is (global_start, global_end, paragraph).
    """
    parts: list[str] = []
    spans: list[tuple[int, int, object]] = []
    pos = 0
    for p in iter_body_paragraphs(doc):
        text = p.text or ""
        start = pos
        end = start + len(text)
        spans.append((start, end, p))
        parts.append(text)
        pos = end + 1  # +1 for join newline
    corpus = "\n".join(parts)
    return corpus, spans


def para_for_offset(
    spans: list[tuple[int, int, object]], global_idx: int
) -> tuple[object, int] | None:
    for start, end, p in spans:
        if start <= global_idx < end or (global_idx == end and start == end):
            return p, global_idx - start
        if start <= global_idx <= end:
            return p, min(global_idx - start, end - start)
    return None


def chunk_ranges(n: int, size: int, overlap: int) -> list[tuple[int, int]]:
    if n <= size:
        return [(0, n)]
    out = []
    i = 0
    while i < n:
        j = min(n, i + size)
        out.append((i, j))
        if j >= n:
            break
        i = max(0, j - overlap)
    return out


def collect_tanakh_refs(corpus: str, base: str) -> list[dict]:
    hits: list[dict] = []
    seen_refs: set[str] = set()
    for lang in ("he", "en"):
        for a, b in chunk_ranges(len(corpus), CHUNK_CHARS, CHUNK_OVERLAP):
            chunk = corpus[a:b]
            if not chunk.strip():
                continue
            result = find_refs(chunk, lang=lang, base=base)
            for hit in tanakh_refs_from_find_result(result):
                hit = dict(hit)
                hit["startChar"] += a
                hit["endChar"] += a
                hits.append(hit)
                seen_refs.add(hit["ref"])
    # Also keep unique refs even if we only need verse text once
    return hits


def verse_cache_get(cache: dict[str, str], tref: str, base: str) -> str:
    if tref not in cache:
        raw = get_text_nikud(tref, base=base)
        cache[tref] = menukad_bli_teamim(raw)
    return cache[tref]


def find_replacements(
    corpus: str,
    spans: list[tuple[int, int, object]],
    refs: list[dict],
    base: str,
    *,
    min_letters: int,
    also_strip_teamim: bool,
) -> list[dict]:
    """Plan replacements: global start/end → pointed text.

    Matching is per paragraph so a full-verse hit in one place does not hide a
    shorter quote (or trop-bearing quote) elsewhere. Partials only apply in
    paragraphs that cite the same ref (avoids random letter collisions).
    """
    cache: dict[str, str] = {}
    planned: list[dict] = []
    occupied: list[tuple[int, int]] = []

    def overlaps(a: int, b: int) -> bool:
        for x, y in occupied:
            if not (b <= x or a >= y):
                return True
        return False

    unique_refs: list[str] = []
    seen: set[str] = set()
    cite_paras: dict[str, set[int]] = {}
    for h in refs:
        if h["ref"] not in seen:
            seen.add(h["ref"])
            unique_refs.append(h["ref"])
        for i, (p_start, p_end, _p) in enumerate(spans):
            if p_start <= h["startChar"] < p_end or p_start < h["endChar"] <= p_end:
                cite_paras.setdefault(h["ref"], set()).add(i)

    def consider(tref: str, pointed: str, g0: int, g1: int, span: str) -> None:
        if not (needs_nikud(span) or (also_strip_teamim and needs_teamim_strip(span))):
            return
        replacement = _pointed_for_letters(pointed, letters_only(span))
        if not replacement or overlaps(g0, g1):
            return
        # No-op only if the visible text is already exactly the replacement
        if span == replacement:
            return
        occupied.append((g0, g1))
        planned.append(
            {
                "ref": tref,
                "start": g0,
                "end": g1,
                "old": span,
                "new": replacement,
                "reason": "missing_nikud" if needs_nikud(span) else "strip_teamim",
            }
        )

    for tref in unique_refs:
        pointed = verse_cache_get(cache, tref, base)
        verse_letters = letters_only(pointed)
        for i, (p_start, p_end, _p) in enumerate(spans):
            para = corpus[p_start:p_end]
            if not letters_only(para):
                continue
            cited_here = i in cite_paras.get(tref, set())
            for start, end, _matched in match_unpointed_in_haystack(
                para, pointed, min_letters=min_letters
            ):
                g0, g1 = p_start + start, p_start + end
                span = corpus[g0:g1]
                span_letters = letters_only(span)
                full = span_letters == verse_letters
                if not full and not cited_here:
                    # Partials only next to an explicit citation for this ref
                    continue
                if not full and len(span_letters) < max(min_letters, min(12, int(0.25 * len(verse_letters)))):
                    continue
                consider(tref, pointed, g0, g1, span)

    planned.sort(key=lambda x: (-(x["end"] - x["start"]), -x["start"]))
    dedup = []
    seen_span: set[tuple[int, int]] = set()
    for p in planned:
        key = (p["start"], p["end"])
        if key in seen_span:
            continue
        seen_span.add(key)
        dedup.append(p)
    return dedup


def _pointed_for_letters(pointed_verse: str, want_letters: str) -> str | None:
    """Extract the pointed substring whose letters equal want_letters."""
    verse = menukad_bli_teamim(pointed_verse)
    letter_idxs: list[int] = []
    letters: list[str] = []
    for i, ch in enumerate(verse):
        if "\u05d0" <= ch <= "\u05ea":
            letter_idxs.append(i)
            letters.append(ch)
    letter_str = "".join(letters)
    idx = letter_str.find(want_letters)
    if idx < 0:
        return None
    c0 = letter_idxs[idx]
    c1 = letter_idxs[idx + len(want_letters) - 1] + 1
    # Include nikud / shin-dot / rafe that trail the last consonant
    while c1 < len(verse) and (
        "\u05b0" <= verse[c1] <= "\u05bd"
        or verse[c1] in "\u05bf\u05c1\u05c2\u05c7"
    ):
        c1 += 1
    if c1 < len(verse) and verse[c1] in "׃׃":
        frag = verse[c0 : c1 + 1]
    else:
        frag = verse[c0:c1]
    return frag.strip()


def apply_replacements(
    doc: Document,
    spans: list[tuple[int, int, object]],
    planned: list[dict],
    revs: RevIds,
) -> int:
    # Apply reverse by global start so earlier offsets stay valid within a para.
    # Cross-paragraph spans are skipped (quotes rarely wrap).
    applied = 0
    for item in sorted(planned, key=lambda x: -x["start"]):
        loc0 = para_for_offset(spans, item["start"])
        loc1 = para_for_offset(spans, item["end"] - 1)
        if loc0 is None or loc1 is None:
            continue
        p0, local_start = loc0
        p1, local_end_last = loc1
        if p0 is not p1:
            print(f"skip cross-paragraph quote for {item['ref']}", file=sys.stderr)
            continue
        local_end = local_end_last + 1
        # Re-find by text in case offsets drifted (should not within reverse apply
        # if each para is independent — but global→local is stable until that para
        # is edited). After one edit in a para, subsequent edits in same para need
        # re-search.
        current = p0.text or ""
        needle = item["old"]
        # Prefer exact local slice if still matches
        if current[local_start:local_end] == needle:
            start, end = local_start, local_end
        else:
            idx = current.find(needle)
            if idx < 0:
                print(f"skip lost span for {item['ref']}: {needle[:40]}…", file=sys.stderr)
                continue
            start, end = idx, idx + len(needle)
        if tracked_replace_span(p0, start, end, item["new"], revs):
            applied += 1
            print(
                f"  {item['ref']}: {item['reason']} → {item['new'][:50]}…"
                if len(item["new"]) > 50
                else f"  {item['ref']}: {item['reason']} → {item['new']}"
            )
    return applied


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--input", "-i", required=True, type=Path, help="Source .docx")
    ap.add_argument(
        "--output",
        "-o",
        type=Path,
        default=None,
        help="Output path (default: overwrite --input in place with tracked changes)",
    )
    ap.add_argument(
        "--no-backup",
        action="store_true",
        help="Do not write <name>-pre-nikud-backup.docx before in-place save",
    )
    ap.add_argument(
        "--base-url",
        default="https://www.sefaria.org",
        help="Sefaria API base (same as SEFARIA_API_BASE_URL in sefaria-mcp)",
    )
    ap.add_argument("--dry-run", action="store_true", help="Plan only; do not write")
    ap.add_argument(
        "--min-letters",
        type=int,
        default=8,
        help="Minimum Hebrew letters to treat as a quote match",
    )
    ap.add_argument(
        "--strip-teamim",
        action="store_true",
        default=True,
        help="Also rewrite already-pointed quotes that still have te'amim (default on)",
    )
    ap.add_argument(
        "--no-strip-teamim",
        action="store_false",
        dest="strip_teamim",
        help="Only fill missing nikud; leave trop alone",
    )
    args = ap.parse_args()

    input_path = args.input.expanduser().resolve()
    if not input_path.exists():
        print(f"missing input: {input_path}", file=sys.stderr)
        return 1
    out = (args.output or input_path).expanduser().resolve()
    if out.parent != input_path.parent and not args.dry_run:
        out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(input_path))
    corpus, spans = build_corpus(doc)
    print(f"Corpus: {len(corpus)} chars, {len(spans)} paragraphs")

    print("Finding Tanakh refs via Sefaria /api/find-refs …")
    refs = collect_tanakh_refs(corpus, args.base_url)
    uniq = sorted({r["ref"] for r in refs})
    print(f"Found {len(refs)} citation hits, {len(uniq)} unique Tanakh refs")
    for r in uniq:
        print(f"  - {r}")

    print("Fetching Tanach with Nikkud and planning replacements …")
    planned = find_replacements(
        corpus,
        spans,
        refs,
        args.base_url,
        min_letters=args.min_letters,
        also_strip_teamim=args.strip_teamim,
    )
    print(f"Planned {len(planned)} replacements")
    for p in planned:
        print(f"  [{p['reason']}] {p['ref']}: {p['old'][:40]!r} → {p['new'][:40]!r}")

    if args.dry_run:
        print("Dry run — no file written")
        return 0

    if not planned:
        print("No changes — file not modified")
        return 0

    if out == input_path and not args.no_backup:
        bak = backup_path_for(input_path)
        if not bak.exists():
            shutil.copy2(input_path, bak)
            print(f"Backup → {bak}")

    enable_track_revisions(doc)
    revs = RevIds(1)
    n = apply_replacements(doc, spans, planned, revs)
    doc.save(str(out))
    print(f"Applied {n} tracked changes → {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
